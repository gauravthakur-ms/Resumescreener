"""Generate standardized LTM resume DOCX by populating the master template.

This module opens the LTM master template (ltm_template.docx) and replaces
candidate-specific content while preserving all branding, images, headers,
footers, logos, page layouts, and section formatting.
"""

import io
import os
import copy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from utils.logger import get_logger

logger = get_logger(__name__)

# Template path
TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "templates")
TEMPLATE_PATH = os.path.join(TEMPLATES_DIR, "ltm_template.docx")


def _clear_paragraph(para):
    """Remove all runs from a paragraph, preserving paragraph-level formatting."""
    for r in para._element.findall(qn('w:r')):
        para._element.remove(r)


def _set_paragraph_text(para, text, bold=False, size=Pt(11), font_name="Calibri"):
    """Set paragraph text with specified formatting, preserving paragraph style."""
    _clear_paragraph(para)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = size
    run.font.name = font_name
    return run


def _replace_textbox_content(body, old_name, old_ps_id, new_name, new_ps_id):
    """Replace text in cover page text boxes (w:txbxContent elements)."""
    txbx_elements = body.findall('.//' + qn('w:txbxContent'))
    for txbx in txbx_elements:
        paras = txbx.findall(qn('w:p'))
        for para in paras:
            texts = [t.text for t in para.findall('.//' + qn('w:t')) if t.text]
            full_text = ''.join(texts)

            # Replace name
            if old_name and old_name in full_text:
                _replace_text_in_xml_para(para, old_name, new_name)
            # Replace PS ID number
            if old_ps_id and old_ps_id in full_text:
                _replace_text_in_xml_para(para, old_ps_id, new_ps_id)


def _replace_text_in_xml_para(para, old_text, new_text):
    """Replace text across runs in an XML paragraph element.

    Handles text split across multiple w:t elements by joining, replacing,
    and redistributing across existing runs.
    """
    t_elements = para.findall('.//' + qn('w:t'))
    if not t_elements:
        return

    # Build full text
    full_text = ''
    for t in t_elements:
        full_text += (t.text or '')

    if old_text not in full_text:
        return

    # Replace in full text
    new_full_text = full_text.replace(old_text, new_text)

    # Redistribute: put all in first, clear the rest
    if t_elements:
        t_elements[0].text = new_full_text
        # Preserve space attribute
        t_elements[0].set(qn('xml:space'), 'preserve')
        for t in t_elements[1:]:
            t.text = ''


def _populate_skills_table(table, skills_data):
    """Populate the Skills Summary table (4x2) with new data."""
    mappings = [
        ("domain", 0),
        ("programming_languages", 1),
        ("tools", 2),
        ("project_overview", 3),
    ]

    for field, row_idx in mappings:
        if row_idx < len(table.rows):
            value_cell = table.rows[row_idx].cells[1]
            # Clear existing content in value cell
            for para in value_cell.paragraphs:
                _clear_paragraph(para)
            if value_cell.paragraphs:
                _set_paragraph_text(
                    value_cell.paragraphs[0],
                    skills_data.get(field, ""),
                    size=Pt(10)
                )


def _set_cell_text(cell, text, bold=False, size=Pt(10)):
    """Set text in a table cell, clearing existing content."""
    # Clear all paragraphs except the first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    if cell.paragraphs:
        para = cell.paragraphs[0]
        _clear_paragraph(para)
        run = para.add_run(text)
        run.bold = bold
        run.font.size = size
        run.font.name = "Calibri"


def _set_cell_bullets(cell, items):
    """Set bullet list in a table cell, clearing existing content."""
    # Clear all existing paragraphs except first
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]._element
        p.getparent().remove(p)

    if not items:
        if cell.paragraphs:
            _clear_paragraph(cell.paragraphs[0])
        return

    # Set first bullet in existing paragraph
    if cell.paragraphs:
        para = cell.paragraphs[0]
        _clear_paragraph(para)
        run = para.add_run(items[0])
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        _apply_bullet_format(para)

    # Add remaining bullets
    for item in items[1:]:
        para = cell.add_paragraph()
        run = para.add_run(item)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        _apply_bullet_format(para)


def _apply_bullet_format(para):
    """Apply bullet formatting to a paragraph."""
    pPr = para._element.get_or_add_pPr()
    # Set list paragraph style
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = parse_xml(f'<w:pStyle {nsdecls("w")} w:val="ListParagraph"/>')
        pPr.insert(0, pStyle)
    else:
        pStyle.set(qn('w:val'), 'ListParagraph')

    # Add bullet numPr if not present
    numPr = pPr.find(qn('w:numPr'))
    if numPr is None:
        numPr = parse_xml(
            f'<w:numPr {nsdecls("w")}>'
            f'  <w:ilvl w:val="0"/>'
            f'  <w:numId w:val="1"/>'
            f'</w:numPr>'
        )
        pPr.append(numPr)


def _remove_table_from_body(table):
    """Remove a table element from the document body."""
    tbl = table._tbl
    tbl.getparent().remove(tbl)


def _remove_paragraph_from_body(para):
    """Remove a paragraph element from the body."""
    p = para._element
    p.getparent().remove(p)


def generate_conversion_docx(data: dict) -> bytes:
    """Generate a standardized LTM resume DOCX from extracted conversion data.

    Opens the LTM master template and replaces candidate-specific content
    while preserving all branding, images, headers, footers, and layout.

    Args:
        data: Dictionary containing all extracted fields (personal, projects, etc.)

    Returns:
        bytes: The generated DOCX file content.
    """
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"LTM template not found at: {TEMPLATE_PATH}")

    doc = Document(TEMPLATE_PATH)
    body = doc.element.body

    personal = data.get("personal", {})
    new_name = personal.get("name", "Candidate Name")
    new_ps_id = personal.get("ps_id", "")

    # === 1. COVER PAGE: Replace Name and PS ID in text boxes ===
    _replace_textbox_content(
        body,
        old_name="Ashwani Kumar Sengar",
        old_ps_id="10735676",
        new_name=new_name,
        new_ps_id=new_ps_id if new_ps_id != "NDATA" else ""
    )

    # === 2. EXPERIENCE SUMMARY: Replace bullet points ===
    experience_summary = data.get("experience_summary", [])
    _replace_experience_summary(doc, experience_summary)

    # === 3. SKILLS SUMMARY: Populate table ===
    skills = data.get("skills_summary", {})
    if doc.tables:
        _populate_skills_table(doc.tables[0], skills)

    # === 4. KEY PROJECTS: Replace project tables ===
    projects = data.get("projects", [])
    _replace_projects(doc, body, projects)

    # === 5. OTHER EXPERIENCE ===
    other_exp = data.get("other_experience", [])
    _replace_other_experience(doc, body, other_exp)

    # === 6. EDUCATIONAL QUALIFICATION ===
    education = data.get("education", [])
    _replace_education(doc, education)

    # === 7. PROFESSIONAL CERTIFICATIONS ===
    certifications = data.get("certifications", [])
    _replace_certifications(doc, certifications)

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _replace_experience_summary(doc, new_bullets):
    """Replace the Experience Summary bullet points."""
    exp_start_idx = None
    exp_end_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "Experience Summary" in text and exp_start_idx is None:
            exp_start_idx = i + 1
        elif exp_start_idx is not None and text in (
            "Skills Summary", "Key Projects", "Other Experience",
            "Educational Qualification", "Professional Certifications"
        ):
            exp_end_idx = i
            break

    if exp_start_idx is None:
        return

    if exp_end_idx is None:
        exp_end_idx = min(exp_start_idx + 15, len(doc.paragraphs))

    # Collect existing bullet paragraphs
    bullet_paras = []
    for i in range(exp_start_idx, exp_end_idx):
        para = doc.paragraphs[i]
        if para.text.strip() and para.style and "List" in para.style.name:
            bullet_paras.append(para)

    if not new_bullets:
        # Remove all bullet paragraphs
        for para in bullet_paras:
            _remove_paragraph_from_body(para)
        return

    # Replace existing bullets and add/remove as needed
    for idx, bullet_text in enumerate(new_bullets):
        if idx < len(bullet_paras):
            _set_paragraph_text(bullet_paras[idx], bullet_text, size=Pt(11))
        else:
            # Clone the format from the first bullet paragraph
            if bullet_paras:
                ref_p = bullet_paras[-1]._element
                new_p = copy.deepcopy(ref_p)
            else:
                new_p = parse_xml(
                    f'<w:p {nsdecls("w")}>'
                    f'<w:pPr><w:pStyle w:val="ListParagraph"/>'
                    f'<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
                    f'</w:pPr></w:p>'
                )
                ref_p = doc.paragraphs[exp_start_idx - 1]._element

            # Clear runs in cloned paragraph
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            ref_p.addnext(new_p)

            # Add text run
            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, None)
            run = new_para.add_run(bullet_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            bullet_paras.append(new_para)

    # Remove extra bullet paragraphs
    for idx in range(len(new_bullets), len(bullet_paras)):
        _remove_paragraph_from_body(bullet_paras[idx])


def _replace_projects(doc, body, projects):
    """Replace Key Projects tables with new project data."""
    if not projects:
        return

    # Find project tables (have "Project Name" in row 0 cell 0)
    project_tables = []
    for table in doc.tables:
        if len(table.rows) >= 4:
            first_cell_text = table.rows[0].cells[0].text.strip()
            if first_cell_text in ("Project Name", "\nProject Name"):
                project_tables.append(table)

    if not project_tables:
        return

    # Keep first project table as format reference
    template_table = project_tables[0]

    # Remove all project tables except the first
    for table in project_tables[1:]:
        _remove_table_from_body(table)

    # Populate the first table with first project
    _populate_project_table(template_table, projects[0])

    # Clone for additional projects
    last_element = template_table._tbl
    for project in projects[1:]:
        # Spacing paragraph
        spacing_p = parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>'
        )
        last_element.addnext(spacing_p)
        last_element = spacing_p

        # Clone template table
        new_tbl = copy.deepcopy(template_table._tbl)
        last_element.addnext(new_tbl)
        last_element = new_tbl

        # Populate cloned table
        from docx.table import Table
        new_table = Table(new_tbl, body)
        _populate_project_table(new_table, project)


def _populate_project_table(table, project_data):
    """Populate a project table (5x4) with project data."""
    rows = table.rows
    if len(rows) < 5:
        return

    # Row 0: [Project Name label] [value] [Team Size label] [value]
    _set_cell_text(rows[0].cells[1], project_data.get("project_name", ""), bold=True)
    try:
        _set_cell_text(rows[0].cells[3], project_data.get("team_size", ""))
    except IndexError:
        pass

    # Row 1: [Start Date label] [value] [End Date label] [value]
    _set_cell_text(rows[1].cells[1], project_data.get("start_date", ""), bold=True)
    try:
        _set_cell_text(rows[1].cells[3], project_data.get("end_date", ""))
    except IndexError:
        pass

    # Row 2: [Project Description label] [value - merged across cols 1-3]
    _set_cell_text(rows[2].cells[1], project_data.get("description", ""))

    # Row 3: [Role & Contribution label] [value - merged, bullet list]
    _set_cell_bullets(rows[3].cells[1], project_data.get("role_contributions", []))

    # Row 4: [Technology & Tools label] [value - merged]
    _set_cell_text(rows[4].cells[1], project_data.get("technologies", ""), bold=True)


def _replace_other_experience(doc, body, other_exp):
    """Replace Other Experience section tables."""
    if not other_exp:
        return

    # Find tables with "Title" in first cell
    other_tables = []
    for table in doc.tables:
        if len(table.rows) >= 4:
            first_cell_text = table.rows[0].cells[0].text.strip()
            if first_cell_text == "Title":
                other_tables.append(table)

    if not other_tables:
        return

    template_table = other_tables[0]

    # Remove extra tables
    for table in other_tables[1:]:
        _remove_table_from_body(table)

    # Populate first table
    _populate_other_exp_table(template_table, other_exp[0])

    # Clone for additional entries
    last_element = template_table._tbl
    for exp in other_exp[1:]:
        spacing_p = parse_xml(
            f'<w:p {nsdecls("w")}><w:pPr><w:spacing w:after="120"/></w:pPr></w:p>'
        )
        last_element.addnext(spacing_p)
        last_element = spacing_p

        new_tbl = copy.deepcopy(template_table._tbl)
        last_element.addnext(new_tbl)
        last_element = new_tbl

        from docx.table import Table
        new_table = Table(new_tbl, body)
        _populate_other_exp_table(new_table, exp)


def _populate_other_exp_table(table, exp_data):
    """Populate an Other Experience table."""
    rows = table.rows
    if len(rows) < 4:
        return

    # Row 0: Title (merged cells 1-3)
    _set_cell_text(rows[0].cells[1], exp_data.get("title", ""))

    # Row 1: Project Name (merged cells 1-3)
    _set_cell_text(rows[1].cells[1], exp_data.get("project_name", ""))

    # Row 2: Start Date, End Date
    _set_cell_text(rows[2].cells[1], exp_data.get("start_date", ""))
    try:
        _set_cell_text(rows[2].cells[3], exp_data.get("end_date", ""))
    except IndexError:
        pass

    # Row 3: Role & Contribution (bullets)
    _set_cell_bullets(rows[3].cells[1], exp_data.get("role_contributions", []))

    # Row 4: Technology & Tools (if present)
    if len(rows) >= 5:
        _set_cell_text(rows[4].cells[1], exp_data.get("technologies", ""), bold=True)


def _replace_education(doc, education):
    """Replace the Education table content."""
    if not education:
        return

    # Find education table
    edu_table = None
    for table in doc.tables:
        if len(table.rows) >= 1:
            first_cell_text = table.rows[0].cells[0].text.strip()
            if "Education" in first_cell_text and "Credentials" in first_cell_text:
                edu_table = table
                break

    if edu_table is None:
        return

    # Build education text
    edu_parts = []
    for edu in education:
        edu_text = edu.get("degree", "")
        if edu.get("year_range"):
            edu_text += f" ({edu['year_range']})"
        if edu.get("institution"):
            edu_text += f" in {edu['institution']}"
        if edu.get("percentage"):
            edu_text += f" with {edu['percentage']}"
        edu_parts.append(edu_text)

    combined = "\n".join(edu_parts)
    _set_cell_text(edu_table.rows[0].cells[1], combined)


def _replace_certifications(doc, certifications):
    """Replace the Professional Certifications bullet points."""
    if not certifications:
        return

    # Find "Professional Certifications" heading
    cert_start_idx = None

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if "Professional Certifications" in text:
            cert_start_idx = i + 1
            break

    if cert_start_idx is None:
        return

    # Find existing cert bullet paragraphs
    cert_paras = []
    for i in range(cert_start_idx, min(cert_start_idx + 20, len(doc.paragraphs))):
        para = doc.paragraphs[i]
        if para.style and "List" in para.style.name:
            if para.text.strip():
                cert_paras.append(para)
        elif para.text.strip():
            break  # Hit next section

    # Build cert texts
    cert_texts = []
    for cert in certifications:
        cert_text = cert.get("name", "")
        validity = cert.get("validity", "")
        if validity:
            cert_text += f" {validity}"
        cert_texts.append(cert_text)

    # Replace existing bullets
    for idx, cert_text in enumerate(cert_texts):
        if idx < len(cert_paras):
            _set_paragraph_text(cert_paras[idx], cert_text, bold=True, size=Pt(11))
        else:
            # Clone from existing cert paragraph
            if cert_paras:
                ref_p = cert_paras[-1]._element
                new_p = copy.deepcopy(ref_p)
            else:
                new_p = parse_xml(
                    f'<w:p {nsdecls("w")}>'
                    f'<w:pPr><w:pStyle w:val="ListParagraph"/>'
                    f'<w:numPr><w:ilvl w:val="0"/><w:numId w:val="1"/></w:numPr>'
                    f'</w:pPr></w:p>'
                )
                ref_p = doc.paragraphs[cert_start_idx - 1]._element

            # Clear runs
            for r in new_p.findall(qn('w:r')):
                new_p.remove(r)
            ref_p.addnext(new_p)

            from docx.text.paragraph import Paragraph
            new_para = Paragraph(new_p, None)
            run = new_para.add_run(cert_text)
            run.font.size = Pt(11)
            run.font.name = "Calibri"
            run.bold = True
            cert_paras.append(new_para)

    # Remove extra cert paragraphs
    for idx in range(len(cert_texts), len(cert_paras)):
        _remove_paragraph_from_body(cert_paras[idx])
